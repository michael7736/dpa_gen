"""
文档解析服务
支持多种格式的文档解析，包括PDF、Word、Markdown等
"""

import asyncio
import io
import logging
import mimetypes
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from abc import ABC, abstractmethod

import aiofiles
from pydantic import BaseModel, Field

# PDF处理
try:
    import pypdf
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word文档处理
try:
    import docx
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Markdown处理
try:
    import markdown
    from markdown.extensions import toc, tables, codehilite
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# 文本处理
import re
from datetime import datetime

from ..models.document import DocumentType, ProcessingStatus
from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentContent(BaseModel):
    """文档内容模型"""
    text: str = Field(..., description="文档文本内容")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="文档元数据")
    structure: Dict[str, Any] = Field(default_factory=dict, description="文档结构信息")
    images: List[Dict[str, Any]] = Field(default_factory=list, description="图片信息")
    tables: List[Dict[str, Any]] = Field(default_factory=list, description="表格信息")


class DocumentSection(BaseModel):
    """文档章节模型"""
    title: str = Field(..., description="章节标题")
    content: str = Field(..., description="章节内容")
    level: int = Field(..., description="章节层级")
    page_number: Optional[int] = Field(None, description="页码")
    section_number: Optional[str] = Field(None, description="章节编号")


class BaseDocumentParser(ABC):
    """文档解析器基类"""
    
    def __init__(self):
        self.logger = get_logger(self.__class__.__name__)
    
    @abstractmethod
    async def parse(self, file_path: Union[str, Path], **kwargs) -> DocumentContent:
        """解析文档"""
        pass
    
    @abstractmethod
    def supports(self, file_path: Union[str, Path]) -> bool:
        """检查是否支持该文件类型"""
        pass
    
    def _extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """提取基础元数据"""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            "filename": path.name,
            "file_size": stat.st_size,
            "file_extension": path.suffix.lower(),
            "mime_type": mimetypes.guess_type(str(path))[0],
            "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }


class PDFDocumentParser(BaseDocumentParser):
    """PDF文档解析器"""
    
    def __init__(self):
        super().__init__()
        if not PDF_AVAILABLE:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """检查是否支持PDF文件"""
        return Path(file_path).suffix.lower() == '.pdf'
    
    async def parse(self, file_path: Union[str, Path], **kwargs) -> DocumentContent:
        """解析PDF文档"""
        self.logger.info(f"开始解析PDF文档: {file_path}")
        
        try:
            # 在线程池中执行PDF解析
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, self._parse_pdf_sync, file_path, kwargs)
            
            self.logger.info(f"PDF文档解析完成: {file_path}")
            return content
            
        except Exception as e:
            self.logger.error(f"PDF文档解析失败: {file_path}, 错误: {str(e)}")
            raise
    
    def _parse_pdf_sync(self, file_path: Union[str, Path], kwargs: Dict[str, Any]) -> DocumentContent:
        """同步解析PDF文档"""
        path = Path(file_path)
        
        with open(path, 'rb') as file:
            reader = PdfReader(file)
            
            # 提取基础元数据
            metadata = self._extract_metadata(file_path)
            
            # 提取PDF元数据
            if reader.metadata:
                pdf_metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "producer": reader.metadata.get("/Producer", ""),
                    "creation_date": reader.metadata.get("/CreationDate", ""),
                    "modification_date": reader.metadata.get("/ModDate", ""),
                }
                metadata.update(pdf_metadata)
            
            # 提取页面信息
            metadata["page_count"] = len(reader.pages)
            metadata["is_encrypted"] = reader.is_encrypted
            
            # 提取文本内容
            text_content = []
            structure = {"sections": [], "pages": []}
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                        structure["pages"].append({
                            "page_number": page_num,
                            "text_length": len(page_text),
                            "has_images": len(page.images) > 0 if hasattr(page, 'images') else False
                        })
                except Exception as e:
                    self.logger.warning(f"无法提取第{page_num}页内容: {str(e)}")
            
            # 合并所有文本
            full_text = "\n\n".join(text_content)
            
            # 提取章节结构
            sections = self._extract_sections_from_text(full_text)
            structure["sections"] = [section.dict() for section in sections]
            
            return DocumentContent(
                text=full_text,
                metadata=metadata,
                structure=structure,
                images=[],  # TODO: 实现图片提取
                tables=[]   # TODO: 实现表格提取
            )
    
    def _extract_sections_from_text(self, text: str) -> List[DocumentSection]:
        """从文本中提取章节结构"""
        sections = []
        
        # 章节标题的正则表达式模式
        patterns = [
            r'^(第[一二三四五六七八九十\d]+章|Chapter\s+\d+|CHAPTER\s+\d+)\s*[：:]\s*(.+)$',
            r'^(\d+\.)\s*(.+)$',
            r'^(\d+\.\d+)\s*(.+)$',
            r'^([一二三四五六七八九十]+、)\s*(.+)$',
            r'^(#{1,6})\s*(.+)$',  # Markdown风格标题
        ]
        
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # 检查是否为章节标题
            is_title = False
            title_level = 0
            title_text = ""
            section_number = ""
            
            for pattern in patterns:
                match = re.match(pattern, line, re.MULTILINE)
                if match:
                    is_title = True
                    section_number = match.group(1)
                    title_text = match.group(2)
                    
                    # 确定标题层级
                    if '章' in section_number or 'Chapter' in section_number:
                        title_level = 1
                    elif re.match(r'^\d+\.$', section_number):
                        title_level = 2
                    elif re.match(r'^\d+\.\d+$', section_number):
                        title_level = 3
                    elif section_number.startswith('#'):
                        title_level = len(section_number)
                    else:
                        title_level = 2
                    
                    break
            
            if is_title:
                # 保存前一个章节
                if current_section and section_content:
                    current_section.content = '\n'.join(section_content).strip()
                    sections.append(current_section)
                
                # 创建新章节
                current_section = DocumentSection(
                    title=title_text,
                    content="",
                    level=title_level,
                    section_number=section_number.strip(),
                    page_number=None  # PDF中难以准确确定页码
                )
                section_content = []
            else:
                # 添加到当前章节内容
                if current_section:
                    section_content.append(line)
        
        # 保存最后一个章节
        if current_section and section_content:
            current_section.content = '\n'.join(section_content).strip()
            sections.append(current_section)
        
        return sections


class WordDocumentParser(BaseDocumentParser):
    """Word文档解析器"""
    
    def __init__(self):
        super().__init__()
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word parsing. Install with: pip install python-docx")
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """检查是否支持Word文件"""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
    
    async def parse(self, file_path: Union[str, Path], **kwargs) -> DocumentContent:
        """解析Word文档"""
        self.logger.info(f"开始解析Word文档: {file_path}")
        
        try:
            # 在线程池中执行Word解析
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, self._parse_word_sync, file_path, kwargs)
            
            self.logger.info(f"Word文档解析完成: {file_path}")
            return content
            
        except Exception as e:
            self.logger.error(f"Word文档解析失败: {file_path}, 错误: {str(e)}")
            raise
    
    def _parse_word_sync(self, file_path: Union[str, Path], kwargs: Dict[str, Any]) -> DocumentContent:
        """同步解析Word文档"""
        path = Path(file_path)
        
        doc = DocxDocument(str(path))
        
        # 提取基础元数据
        metadata = self._extract_metadata(file_path)
        
        # 提取Word文档属性
        core_props = doc.core_properties
        word_metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "comments": core_props.comments or "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision,
            "created": core_props.created.isoformat() if core_props.created else "",
            "modified": core_props.modified.isoformat() if core_props.modified else "",
        }
        metadata.update(word_metadata)
        
        # 提取文本内容和结构
        text_content = []
        sections = []
        tables = []
        images = []
        
        current_section = None
        section_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检查是否为标题
            if paragraph.style.name.startswith('Heading'):
                # 保存前一个章节
                if current_section and section_content:
                    current_section.content = '\n'.join(section_content).strip()
                    sections.append(current_section)
                
                # 创建新章节
                level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                current_section = DocumentSection(
                    title=text,
                    content="",
                    level=level,
                    section_number="",
                    page_number=None
                )
                section_content = []
            else:
                # 添加到内容
                text_content.append(text)
                if current_section:
                    section_content.append(text)
        
        # 保存最后一个章节
        if current_section and section_content:
            current_section.content = '\n'.join(section_content).strip()
            sections.append(current_section)
        
        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            if table_data:
                tables.append({
                    "data": table_data,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0
                })
        
        # 合并所有文本
        full_text = '\n'.join(text_content)
        
        structure = {
            "sections": [section.dict() for section in sections],
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(tables),
            "image_count": len(images)
        }
        
        return DocumentContent(
            text=full_text,
            metadata=metadata,
            structure=structure,
            images=images,
            tables=tables
        )


class MarkdownDocumentParser(BaseDocumentParser):
    """Markdown文档解析器"""
    
    def __init__(self):
        super().__init__()
        if not MARKDOWN_AVAILABLE:
            raise ImportError("markdown is required for Markdown parsing. Install with: pip install markdown")
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """检查是否支持Markdown文件"""
        return Path(file_path).suffix.lower() in ['.md', '.markdown']
    
    async def parse(self, file_path: Union[str, Path], **kwargs) -> DocumentContent:
        """解析Markdown文档"""
        self.logger.info(f"开始解析Markdown文档: {file_path}")
        
        try:
            # 异步读取文件
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            # 解析Markdown
            result = await self._parse_markdown_content(content, file_path)
            
            self.logger.info(f"Markdown文档解析完成: {file_path}")
            return result
            
        except Exception as e:
            self.logger.error(f"Markdown文档解析失败: {file_path}, 错误: {str(e)}")
            raise
    
    async def _parse_markdown_content(self, content: str, file_path: Union[str, Path]) -> DocumentContent:
        """解析Markdown内容"""
        # 提取基础元数据
        metadata = self._extract_metadata(file_path)
        
        # 配置Markdown解析器
        md = markdown.Markdown(
            extensions=[
                'toc',
                'tables',
                'codehilite',
                'fenced_code',
                'footnotes',
                'attr_list',
                'def_list'
            ],
            extension_configs={
                'toc': {
                    'permalink': True,
                    'title': 'Table of Contents'
                }
            }
        )
        
        # 转换为HTML（用于结构分析）
        html = md.convert(content)
        
        # 提取目录结构
        toc = getattr(md, 'toc_tokens', [])
        
        # 提取章节结构
        sections = self._extract_markdown_sections(content)
        
        # 提取表格
        tables = self._extract_markdown_tables(content)
        
        # 提取图片
        images = self._extract_markdown_images(content)
        
        # 提取纯文本内容
        plain_text = self._markdown_to_text(content)
        
        structure = {
            "sections": [section.dict() for section in sections],
            "toc": toc,
            "heading_count": len([s for s in sections if s.level <= 3]),
            "table_count": len(tables),
            "image_count": len(images),
            "code_block_count": len(re.findall(r'```[\s\S]*?```', content))
        }
        
        # 更新元数据
        metadata.update({
            "word_count": len(plain_text.split()),
            "character_count": len(plain_text),
            "line_count": len(content.split('\n'))
        })
        
        return DocumentContent(
            text=plain_text,
            metadata=metadata,
            structure=structure,
            images=images,
            tables=tables
        )
    
    def _extract_markdown_sections(self, content: str) -> List[DocumentSection]:
        """提取Markdown章节结构"""
        sections = []
        lines = content.split('\n')
        
        current_section = None
        section_content = []
        
        for line in lines:
            # 检查是否为标题
            if line.startswith('#'):
                # 保存前一个章节
                if current_section and section_content:
                    current_section.content = '\n'.join(section_content).strip()
                    sections.append(current_section)
                
                # 解析新标题
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                
                current_section = DocumentSection(
                    title=title,
                    content="",
                    level=level,
                    section_number="",
                    page_number=None
                )
                section_content = []
            else:
                # 添加到当前章节内容
                if current_section:
                    section_content.append(line)
        
        # 保存最后一个章节
        if current_section and section_content:
            current_section.content = '\n'.join(section_content).strip()
            sections.append(current_section)
        
        return sections
    
    def _extract_markdown_tables(self, content: str) -> List[Dict[str, Any]]:
        """提取Markdown表格"""
        tables = []
        
        # 匹配Markdown表格
        table_pattern = r'(\|.*\|.*\n\|.*[-:]+.*\|.*\n(?:\|.*\|.*\n)*)'
        
        for match in re.finditer(table_pattern, content, re.MULTILINE):
            table_text = match.group(1)
            lines = table_text.strip().split('\n')
            
            if len(lines) >= 2:
                # 解析表头
                header = [cell.strip() for cell in lines[0].split('|')[1:-1]]
                
                # 解析数据行
                data_rows = []
                for line in lines[2:]:  # 跳过分隔符行
                    row = [cell.strip() for cell in line.split('|')[1:-1]]
                    if row:
                        data_rows.append(row)
                
                tables.append({
                    "header": header,
                    "data": data_rows,
                    "rows": len(data_rows),
                    "columns": len(header)
                })
        
        return tables
    
    def _extract_markdown_images(self, content: str) -> List[Dict[str, Any]]:
        """提取Markdown图片"""
        images = []
        
        # 匹配Markdown图片语法
        image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        
        for match in re.finditer(image_pattern, content):
            alt_text = match.group(1)
            url = match.group(2)
            
            images.append({
                "alt_text": alt_text,
                "url": url,
                "type": "markdown_image"
            })
        
        return images
    
    def _markdown_to_text(self, content: str) -> str:
        """将Markdown转换为纯文本"""
        # 移除Markdown语法
        text = content
        
        # 移除标题标记
        text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
        
        # 移除链接，保留文本
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # 移除图片
        text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
        
        # 移除粗体和斜体标记
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # 移除代码块
        text = re.sub(r'```[\s\S]*?```', '', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # 移除表格分隔符
        text = re.sub(r'\|.*[-:]+.*\|', '', text)
        
        # 清理多余的空行
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()


class TextDocumentParser(BaseDocumentParser):
    """纯文本文档解析器"""
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """检查是否支持文本文件"""
        return Path(file_path).suffix.lower() in ['.txt', '.text']
    
    async def parse(self, file_path: Union[str, Path], **kwargs) -> DocumentContent:
        """解析文本文档"""
        self.logger.info(f"开始解析文本文档: {file_path}")
        
        try:
            # 异步读取文件
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            # 提取基础元数据
            metadata = self._extract_metadata(file_path)
            metadata.update({
                "word_count": len(content.split()),
                "character_count": len(content),
                "line_count": len(content.split('\n'))
            })
            
            # 简单的段落分割
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            structure = {
                "paragraph_count": len(paragraphs),
                "sections": []
            }
            
            self.logger.info(f"文本文档解析完成: {file_path}")
            
            return DocumentContent(
                text=content,
                metadata=metadata,
                structure=structure,
                images=[],
                tables=[]
            )
            
        except Exception as e:
            self.logger.error(f"文本文档解析失败: {file_path}, 错误: {str(e)}")
            raise


class DocumentParserFactory:
    """文档解析器工厂"""
    
    def __init__(self):
        self.parsers = []
        self.logger = get_logger(__name__)
        
        # 注册默认解析器
        self._register_default_parsers()
    
    def _register_default_parsers(self):
        """注册默认解析器"""
        try:
            if PDF_AVAILABLE:
                self.register_parser(PDFDocumentParser())
        except ImportError:
            self.logger.warning("PDF解析器不可用，请安装pypdf")
        
        try:
            if DOCX_AVAILABLE:
                self.register_parser(WordDocumentParser())
        except ImportError:
            self.logger.warning("Word解析器不可用，请安装python-docx")
        
        try:
            if MARKDOWN_AVAILABLE:
                self.register_parser(MarkdownDocumentParser())
        except ImportError:
            self.logger.warning("Markdown解析器不可用，请安装markdown")
        
        # 文本解析器总是可用
        self.register_parser(TextDocumentParser())
    
    def register_parser(self, parser: BaseDocumentParser):
        """注册解析器"""
        self.parsers.append(parser)
        self.logger.info(f"已注册解析器: {parser.__class__.__name__}")
    
    def get_parser(self, file_path: Union[str, Path]) -> Optional[BaseDocumentParser]:
        """获取适合的解析器"""
        for parser in self.parsers:
            if parser.supports(file_path):
                return parser
        return None
    
    async def parse_document(self, file_path: Union[str, Path], **kwargs) -> DocumentContent:
        """解析文档"""
        parser = self.get_parser(file_path)
        if not parser:
            raise ValueError(f"不支持的文档类型: {file_path}")
        
        return await parser.parse(file_path, **kwargs)
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        extensions = []
        test_files = [
            Path("test.pdf"),
            Path("test.docx"),
            Path("test.doc"),
            Path("test.md"),
            Path("test.markdown"),
            Path("test.txt"),
            Path("test.text")
        ]
        
        for test_file in test_files:
            if self.get_parser(test_file):
                extensions.append(test_file.suffix.lower())
        
        return list(set(extensions))


# 全局解析器工厂实例
document_parser_factory = DocumentParserFactory()


# 便捷函数
async def parse_document(file_path: Union[str, Path], **kwargs) -> DocumentContent:
    """解析文档的便捷函数"""
    return await document_parser_factory.parse_document(file_path, **kwargs)


def get_supported_extensions() -> List[str]:
    """获取支持的文件扩展名"""
    return document_parser_factory.get_supported_extensions()


def is_supported_document(file_path: Union[str, Path]) -> bool:
    """检查是否支持该文档类型"""
    return document_parser_factory.get_parser(file_path) is not None